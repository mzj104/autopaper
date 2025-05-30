import os.path

from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
import sys
from launch import myluanch, path
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from getpic import download_pic
from read import load

def set_font(run, font_name, size_pt, bold=False, color_rgb=(0, 0, 0)):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color_rgb)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    set_font(run, font_name="宋体", size_pt=16, bold=True)  # 三号
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, font_name="宋体", size_pt=14, bold=True)  # 四号
    return p

def add_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, font_name="宋体", size_pt=12, bold=False)  # 小四
    p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进 2 字符
    return p

def add_newline_every_two_periods(text):
    count = 0
    result = ""
    for char in text:
        result += char
        if char == '。':
            count += 1
            if count % 2 == 0:
                result += '\n\n'
    return result

def gen():
    port = 8123
    proc = myluanch(port, 'https://chatgpt.com/')

    chrome_driver_path = path + r'chromedriver.exe'
    chrome_options = Options()
    chrome_options.add_experimental_option("debuggerAddress", f"127.0.0.1:{port}")  # 连接已经打开的浏览器
    service = Service(executable_path=chrome_driver_path)
    driver = webdriver.Chrome(service=service, options=chrome_options)
    wait = WebDriverWait(driver, 20)

    doc = Document()
    data = load('data.json')
    lis = []
    title = ""
    for i in data['大纲']:
        tmp = data['大纲'][i]
        lis = tmp
        title = i
        break


    cnt = 1
    for i in lis:
        text_area = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, 'div[contenteditable="true"]')))
        prompt =f'你现在需要写营销号文章，总的标题是：{title}，现在写第{cnt}部分：{i["标题"]}，内容：{i["内容"]}，内容要充实，尽可能多说一些话，减少ai表达，不能使用更细分的标题或序号，用文字把他们串联。不要使用“首先”“其次”“总的来说”这种一看就是ai的表达。'
        if title != '总结' and '结语':
            prompt += '不要加入总结性的语句，因为这只是一个小段，加入会显得很假。'
        text_area.send_keys(prompt)
        time.sleep(1)
        input_value = text_area.text
        print(f"输入框中的内容是: {input_value}")
        send_button = wait.until(EC.element_to_be_clickable((By.ID, "composer-submit-button")))
        send_button.click()
        print("发送按钮已点击")
        last_answer = ""
        timeout = 5
        start_time = time.time()

        while True:
            answers = driver.find_elements(By.CSS_SELECTOR, 'div[data-message-author-role="assistant"]')
            if answers:
                new_answer = answers[-1].text.strip()
                if new_answer != last_answer and new_answer != "":
                    with open('./output/out'+str(cnt)+'.txt','w', encoding='utf-8') as f:
                        f.write(new_answer)
                    last_answer = new_answer
                    start_time = time.time()

            if time.time() - start_time > timeout:
                print("⏰ 超时退出监听")
                break

            time.sleep(1)  # 每秒检查一次
        cnt += 1

    all = ""
    pic_cnt = 0
    for i in range(1, cnt):
        with open('./output/out' + str(i) + '.txt', 'r', encoding='utf-8') as f:
            tmp = f.read()
            tmp = tmp.replace('\n', '')
            tmp = add_newline_every_two_periods(tmp)
            t = lis[i-1]['标题']
            add_subtitle(doc, t)
            ins = 0
            split = tmp.split('\n\n')
            for k in split:
                add_paragraph(doc, k)
                ins += 1
                if ins == 4:
                    pic_path = f'./images/image_{pic_cnt}.jpg'
                    while not os.path.exists(pic_path):
                        pic_cnt += 1
                        pic_path = f'./images/image_{pic_cnt}.jpg'

                    doc.add_picture(f'./images/image_{pic_cnt}.jpg', height=Cm(10))
                    pic_cnt += 1
                    ins = 0


    doc.save("./words/"+ title + ".docx")
    proc.terminate()