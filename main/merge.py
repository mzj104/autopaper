from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json

def add_newline_every_two_periods(text):
    count = 0
    result = ""
    for char in text:
        result += char
        if char == '。':
            count += 1
            if count % 2 == 0:
                result += '\n\n'
    return result


def set_font(run, font_name, size_pt, bold=False, color_rgb=(0, 0, 0)):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color_rgb)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    set_font(run, font_name="宋体", size_pt=16, bold=True)  # 三号
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, font_name="宋体", size_pt=14, bold=True)  # 四号
    return p

def add_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, font_name="宋体", size_pt=12, bold=False)  # 小四
    p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进 2 字符
    return p

doc = Document()
with open('data.json', 'r', encoding='utf-8') as f:
    data = json.load(f)
lis = []
title = ""
for i in data['大纲']:
    tmp = data['大纲'][i]
    lis = tmp
    title = i
    add_title(doc, title)
    break

cnt = 6
all = ""
for i in range(1, cnt + 1):
    with open('out' + str(i) + '.txt', 'r', encoding='utf-8') as f:
        tmp = f.read()
        tmp = tmp.replace('\n', '')
        tmp = add_newline_every_two_periods(tmp)
        t = lis[i-1]['标题']
        add_subtitle(doc, t)
        split = tmp.split('\n\n')
        for k in split:
            add_paragraph(doc, k)


doc.save("formatted_doc_with_indent.docx")
